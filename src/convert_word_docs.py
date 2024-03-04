import os
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Inches
import re


def convert_word_doc(docx_path, images_folder, output_folder):
    # Ensure the output folder exists
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    document = Document(docx_path)
    relevant_paragraphs = get_relevant_paragraphs(document)
    slides_info = get_slides_info(document, images_folder)
    header_1_paragraphs = get_header_1_paragraphs(document)

    # check that the number of relevant_paragrphs is equal to the sum of the number of slides_info and the number of header_1_paragraphs
    # the number of slides_info should be the same as the number of pictures that we have to add to the new document
    # the header_1_paragraphs are the titles of sections and therefore have no specific slide
    # all other paragraphs are represent the titles of the slides which have a corresponding slide_info
    if len(relevant_paragraphs) != len(slides_info) + len(header_1_paragraphs):
        # raise ValueError(
        print(
            "The number of slide titles is not equal to the sum of the number of relevant tables and the number of header 1 paragraphs\n"
            + f"({len(relevant_paragraphs)} != {len(slides_info)} + {len(header_1_paragraphs)})\n"
            + f"for document {docx_path}"
        )

    # create a new document
    document = Document()

    levels = {
        "Titolo 11": 1,
        "Titolo 21": 2,
        "Titolo 31": 3,
        "Heading 1": 1,
        "Heading 2": 2,
        "Heading 3": 3,
        "Normal": 0,
    }

    # loop through the relevant_paragraphs. For each paragraph:
    # - if the paragraph has style "Titolo 11" then add the paragraph to the new document with heading 1 and move to the next paragraph
    # - if the paragraph has style not equal to "Titolo 11" then
    # - - add the paragraph to the new document
    # - - add the corresponding picture to the new document
    # - - add the content of the paragraph to the new document
    # - - add a page break

    content_index = 0
    paragraph_index = 0
    for paragraph in relevant_paragraphs:
        _paragraph = paragraph["paragraph"]
        # break if the content_index is greater than the length of slides_info and print a warning
        if content_index >= len(slides_info):
            # print the text of the paragraphs that are not added to the new document
            for p in relevant_paragraphs[paragraph_index:]:
                print(f">>>>>>>>>>>>>>>>>paragraph not added: {p['text']}")
            break
        slide_info = slides_info[content_index]
        image_path = slide_info["image_path"]
        slide_notes = slide_info["slide_notes"]
        if _paragraph.style.name == "Titolo 11":
            document.add_heading(_paragraph.text, level=levels[_paragraph.style.name])
        else:
            document.add_heading(_paragraph.text, level=levels[_paragraph.style.name])
            try:
                document.add_picture(image_path, width=Inches(6))
            except:
                print(
                    f">>>>>>>>>>>>>>>>>>>>>>> Image not found: {image_path} for paragraph {_paragraph.text}"
                )
            document.add_paragraph(slide_notes)
            document.add_page_break()
            content_index += 1
        paragraph_index += 1

    # save the document
    # read the original document name
    doc_name = os.path.basename(docx_path)
    output_doc_name = os.path.join(output_folder, f"{doc_name}_converted.docx")
    document.save(output_doc_name)
    return output_doc_name


def get_relevant_paragraphs(document):
    relevant_paragraphs = []
    for paragraph in document.paragraphs:
        # append to relevant_paragraphs list the paragraph ignoring the ones that are empty
        if paragraph.text:
            # use a dictionary for debugging purposes since it is easier to control the text of the paragraph while debugging
            relevant_paragraphs.append({"text": paragraph.text, "paragraph": paragraph})

    # check that the last paragraphs are what we expect them to be before removing them
    check_last_2_paragraphs(
        relevant_paragraphs[-1]["text"], relevant_paragraphs[-2]["text"]
    )
    # remove the last two paragraphs
    relevant_paragraphs = relevant_paragraphs[:-2]

    return relevant_paragraphs


def get_header_1_paragraphs(document):
    paragraphs = document.paragraphs
    # remove the paragraphs that are not header 1
    header_1_paragraphs = [
        p for p in paragraphs if p.style.name in ["Titolo 11", "Heading 1"]
    ]

    # check that the last paragraphs are what we expect them to be before removing them
    check_last_2_paragraphs(header_1_paragraphs[-1].text, header_1_paragraphs[-2].text)
    # remove the last two paragraphs
    header_1_paragraphs = header_1_paragraphs[:-2]

    return header_1_paragraphs


def check_last_2_paragraphs(last_text, second_last_text):
    # check that the last paragraph is "Menu" and that the second to last paragraph is "Player Properties"
    if last_text != "Menu":
        raise ValueError("The last paragraph should be 'Menu'")
    if second_last_text != "Player Properties":
        raise ValueError("The second to last paragraph should be 'Player Properties'")


def get_slides_info(document, image_folder):
    # extract the tables that contain the slide notes
    # these tables will be used to create the content of the slides
    tables_with_slide_notes = get_tables_with_slide_notes(document)

    # convert the tables to a list of dictionaries with the relevant information for each slide
    slides_info = to_slides_info(tables_with_slide_notes, image_folder)

    return slides_info


def get_tables_with_slide_notes(document):
    tables = document.tables
    # info about slides are contained in some of the tables of the document. Such tables can be identified as follows
    # - if the first cell in the first row is "Slide ID 🔒" and the table length is = 2, then the table containing the
    #   slide notes is the next table
    # - if the first cell in the first row is "Layer ID 🔒" and the table length is = 2, then the table containing the
    #   slide notes is the next table
    # - if the first cell in the first row is "Slide ID 🔒" but the length of the table is greater than 2, then the table
    #   itself contains the slide notes (this is probably an edge case - it has been encountered in the first table of C_Sales_Network_Tech.docx)
    tables_with_slide_notes = []
    for i in range(len(tables)):
        table = tables[i]
        if table.cell(0, 0).text in ["Slide ID 🔒", "Layer ID 🔒"]:
            if len(table.rows) == 2:
                next_table = tables[i + 1]
                tables_with_slide_notes.append(next_table)
                i += 1
            elif len(table.rows) > 2:
                tables_with_slide_notes.append(table)
            else:
                error_text = (
                    f'Table {i} which is a "Slide ID 🔒" table - has length < 2'
                )
                raise ValueError(error_text)

    return tables_with_slide_notes


def to_slides_info(tables_with_slide_notes, image_folder):
    # for each table in tables_with_slide_notes, create a dictionary with the following keys:
    # - slide_notes: the content to be shown on the slide
    # - image_path: the path of the image that has to be added to the slide
    slides_info = []
    for i in range(len(tables_with_slide_notes)):
        table = tables_with_slide_notes[i]
        slide_info = {"slide_notes": "", "image_path": ""}
        # the slide notes are identified in the following way:
        # - ignore the first 2 rows if the first cell of the second row is "Slide name" or just the first row otherwise
        #   (the latter is the case for "Slide layers", i.e. the slides identified by "Layer ID 🔒" tables)
        # - if the first cell of the last row is "Slide Notes" then the slide notes are the value of the third cell of the last row
        # - otherwise we have to take one row for each value in the second cell and concatenate the values of the third cell for such rows
        #   (the reason we take just one row for each value in the second cell is that the rows with the same value in the second cell
        #   have the same value in the third cell and if we do not take just one row for each value in the second cell we would have
        #   the same content repeated many times)
        start = 2 if table.rows[1].cells[1].text == "Slide name" else 1
        table_rows = table.rows[start:]
        last_row = table_rows[-1]
        if last_row.cells[1].text == "Slide Notes":
            slide_info["slide_notes"] = last_row.cells[2].text
        else:
            slide_info["slide_notes"] = ""
            previous_value_of_cell_1 = ""
            for row in table_rows:
                if row.cells[1].text != previous_value_of_cell_1:
                    slide_info["slide_notes"] += row.cells[2].text + "\n"
                    previous_value_of_cell_1 = row.cells[1].text
        slide_info["slide_notes"] = format_string(slide_info["slide_notes"])

        # the image path is the path of the image that has to be added to the slide
        # the image path is obtained joining the image folder path with the fixed "image" string followed by
        # the index of the table in the list of tables_with_slide_notes
        slide_info["image_path"] = os.path.join(
            image_folder, "word", "media", f"image{i + 1}.png"
        )

        slides_info.append(slide_info)

    return slides_info


def format_string(s):
    formattedString = s
    if formattedString.startswith("1"):
        formattedString = formattedString[1:]
    # replace the occurrences of the strings "\ni" where i is a number with a new line character
    formattedString = re.sub("\\n\d+", "\n", formattedString)
    formattedString = bullet_string(formattedString)
    return formattedString


def bullet_string(s):
    lines = s.split("\n")
    add_dash = False
    for i in range(len(lines)):
        l = lines[i].strip()
        if l.endswith(":"):
            add_dash = True
        elif l and l[0].isupper():
            add_dash = False
        if add_dash and l:
            lines[i] = "- " + lines[i]
    return "\n".join(lines)
